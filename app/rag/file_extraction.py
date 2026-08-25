"""
File-to-text extraction for `POST /v1/documents/upload`.

Each format's parser is imported lazily, inside the function that needs it
— so the service only requires the dependency matching the file types you
actually plan to accept. Uploading a .pdf without `pypdf` installed fails
with a clear, typed, actionable error instead of an ImportError buried in a
stack trace. .txt/.md need no dependency at all and always work.
"""
from __future__ import annotations

import io
import os

from app.core.exceptions import RagServiceError

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class UnsupportedFileTypeError(RagServiceError):
    code = "unsupported_file_type"
    status_code = 415


class FileExtractionError(RagServiceError):
    code = "file_extraction_failed"
    status_code = 422


def extract_text(filename: str, data: bytes) -> str:
    """Route to the right extractor based on file extension and return
    plain text ready for chunking. Raises a typed RagServiceError on any
    failure so the API returns a structured error, never a raw traceback."""
    ext = os.path.splitext(filename.lower())[1]

    if ext in (".txt", ".md"):
        return _extract_plain_text(data)
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)

    raise UnsupportedFileTypeError(
        f"Unsupported file type '{ext or '(none)'}'. "
        f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
        details={"filename": filename, "extension": ext},
    )


def _extract_plain_text(data: bytes) -> str:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as e:
        raise FileExtractionError("File is not valid UTF-8 text.") from e
    if not text.strip():
        raise FileExtractionError("File is empty after decoding.")
    return text


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise FileExtractionError(
            "PDF extraction requires the 'pypdf' package. Install it with: pip install pypdf"
        ) from e

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages).strip()
    except Exception as e:  # noqa: BLE001 - any parse failure becomes a typed error
        raise FileExtractionError(f"Failed to parse PDF: {e}") from e

    if not text:
        raise FileExtractionError(
            "No extractable text found in this PDF — it may be a scanned image "
            "without an OCR text layer."
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise FileExtractionError(
            "DOCX extraction requires the 'python-docx' package. Install it with: pip install python-docx"
        ) from e

    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(p for p in parts if p.strip()).strip()
    except Exception as e:  # noqa: BLE001
        raise FileExtractionError(f"Failed to parse DOCX: {e}") from e

    if not text:
        raise FileExtractionError("No extractable text found in this DOCX file.")
    return text
